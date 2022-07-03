from docx import Document
from docx.shared import Inches

def generate_dummy_docx_file(firstName, lastName, SID,  statsBasedDiagnostic, anamnesisGen, berteilung):


    document = Document()

    document.add_heading(firstName + ' ' + lastName + ' ' + SID, 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Diagnostic: ' + statsBasedDiagnostic, level=1)


    document.add_heading('Anamnesis', level=1)
    document.add_paragraph(str(anamnesisGen.strip()), style='Intense Quote')

    document.add_heading('Berteilung', level=1)
    document.add_paragraph(str(berteilung.strip()), style='Intense Quote')

    document.add_page_break()

    document.save('report.docx')


if __name__ == "__main__":
    generate_dummy_docx_file()